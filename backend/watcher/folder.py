import os
import sys
import logging
from pathlib import Path
from datetime import datetime
import hashlib
from dataclasses import dataclass
import threading
import time
from typing import Dict, Any, Optional, Set, TypeVar
import requests
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import tkinter as tk
from tkinter import filedialog
from queue import Queue
import mimetypes
import chardet
import logging.handlers

# Text extraction imports
import pdfplumber
from docx import Document
import magic
from PIL import Image
import pytesseract
from dotenv import load_dotenv


def setup_logging():
    """Configure comprehensive logging system with both file and console output."""
    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG)

    # Detailed formatter for debugging
    formatter = logging.Formatter(
        "[%(asctime)s] %(levelname)s [%(name)s:%(funcName)s:%(lineno)d] - %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    # Console Handler - INFO level
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(formatter)

    # File Handler - DEBUG level with rotation
    file_handler = logging.handlers.RotatingFileHandler(
        "folder_watcher.log",
        maxBytes=5 * 1024 * 1024,  # 5MB
        backupCount=5,
        encoding="utf-8",
    )
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(formatter)

    # Remove existing handlers
    for handler in logger.handlers[:]:
        logger.removeHandler(handler)

    logger.addHandler(console_handler)
    logger.addHandler(file_handler)

    return logger


logger = setup_logging()


# dict helper function
T = TypeVar("T")


class JsonHelper:
    """Helper class for safely handling JSON data."""

    @staticmethod
    def get_value(data: dict, key: str, default: Any = None) -> Any:
        """
        Safely get a value from a dictionary with default value.

        Args:
            data: The dictionary to get the value from
            key: The key to look for
            default: Default value to return if key doesn't exist or value is invalid

        Returns:
            The value if it exists, default value otherwise
        """
        try:
            if not isinstance(data, dict):
                logger.warning(f"Input data is not a dictionary: {type(data)}")
                return default

            if key not in data:
                logger.debug(f"Key '{key}' not found in data, using default: {default}")
                return default

            value = data[key]
            if value is None:
                logger.debug(f"Value for key '{key}' is None, using default: {default}")
                return default

            return value

        except Exception as e:
            logger.error(f"Error getting value for key '{key}': {str(e)}")
            return default


@dataclass
class ExtractionResult:
    """Container for extraction results."""

    content: str
    metadata: Dict[str, Any]
    error: Optional[str] = None
    quality_score: float = 1.0


# Supported file extensions
SUPPORTED_EXTENSIONS: Set[str] = {".pdf", ".doc", ".docx", ".txt"}


class DocumentProcessor:
    """Handles document processing and text extraction."""

    def __init__(self):
        # Initialize mime types
        mimetypes.init()
        # Configure minimum content quality thresholds
        self.min_content_length = 50
        self.min_quality_score = 0.3

    def process_file(self, file_path: str) -> ExtractionResult:
        """Process a file and extract text with metadata."""
        try:
            # Detect file type
            mime_type = magic.from_file(file_path, mime=True)

            # Get basic metadata first
            metadata = self._generate_base_metadata(file_path)

            # Extract text based on file type
            if mime_type == "application/pdf":
                result = self._extract_pdf(file_path)
            elif mime_type in [
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ]:
                result = self._extract_doc(file_path)
            elif mime_type.startswith("text/"):
                result = self._extract_text(file_path)
            elif mime_type.startswith("image/"):
                result = self._extract_image(file_path)
            else:
                # Fallback to textract for other formats
                result = self._extract_fallback(file_path)

            # Merge metadata
            metadata.update(result.metadata)

            # Post-process and validate content
            processed_content = self._post_process_content(result.content)
            quality_score = self._calculate_quality_score(
                processed_content, result.metadata
            )

            return ExtractionResult(
                content=processed_content,
                metadata=metadata,
                quality_score=quality_score,
            )

        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return ExtractionResult(
                content="",
                metadata=(
                    metadata
                    if "metadata" in locals()
                    else self._generate_base_metadata(file_path)
                ),
                error=str(e),
                quality_score=0.0,
            )

    def _extract_pdf(self, file_path: str) -> ExtractionResult:
        """Extract text from PDF files."""
        try:
            extracted_text = []
            metadata = {}

            with pdfplumber.open(file_path) as pdf:
                metadata.update(
                    {
                        "page_count": len(pdf.pages),
                        "pdf_info": pdf.metadata,
                    }
                )

                # Extract text from each page with layout preservation
                for i, page in enumerate(pdf.pages):
                    # Extract text with layout analysis
                    text = page.extract_text(layout=True)
                    if text:
                        extracted_text.append(text)

                    # Extract tables if present
                    tables = page.extract_tables()
                    if tables:
                        metadata[f"page_{i+1}_tables"] = len(tables)
                        # Convert tables to text representation
                        for table in tables:
                            table_text = "\n".join(
                                [" | ".join(row) for row in table if any(row)]
                            )
                            extracted_text.append(table_text)

            content = "\n\n".join(extracted_text)

            return ExtractionResult(
                content=content,
                metadata=metadata,
                quality_score=self._calculate_quality_score(content, metadata),
            )

        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise

    def _extract_doc(self, file_path: str) -> ExtractionResult:
        """Extract text from DOC/DOCX files."""
        try:
            doc = Document(file_path)
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
                "has_tables": bool(len(doc.tables)),
            }

            # Extract text with structure preservation
            content_parts = []

            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)

            # Process tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    content_parts.append("\n".join(table_text))

            content = "\n\n".join(content_parts)

            return ExtractionResult(
                content=content,
                metadata=metadata,
                quality_score=self._calculate_quality_score(content, metadata),
            )

        except Exception as e:
            logger.error(f"DOC extraction error: {str(e)}")
            raise

    def _extract_text(self, file_path: str) -> ExtractionResult:
        """Extract text from plain text files with encoding detection."""
        try:
            # Detect encoding
            with open(file_path, "rb") as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result["encoding"]
                confidence = result["confidence"]

            # Read with detected encoding
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()

            metadata = {
                "encoding": encoding,
                "encoding_confidence": confidence,
                "line_count": content.count("\n") + 1,
            }

            return ExtractionResult(
                content=content, metadata=metadata, quality_score=confidence
            )

        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            raise

    def _extract_image(self, file_path: str) -> ExtractionResult:
        """Extract text from images using OCR."""
        try:
            with Image.open(file_path) as img:
                metadata = {
                    "image_format": img.format,
                    "image_size": img.size,
                    "color_mode": img.mode,
                }

                # Perform OCR
                content = pytesseract.image_to_string(img)
                confidence = float(
                    pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)[
                        "conf"
                    ][0]
                )

                return ExtractionResult(
                    content=content, metadata=metadata, quality_score=confidence / 100.0
                )

        except Exception as e:
            logger.error(f"Image extraction error: {str(e)}")
            raise

    def _extract_fallback(self, file_path: str) -> ExtractionResult:
        """Fallback extraction using textract."""
        pass
        # try:
        #     content = textract.process(file_path).decode("utf-8")
        #     metadata = {"extraction_method": "textract"}

        #     return ExtractionResult(content=content, metadata=metadata)

        # except Exception as e:
        #     logger.error(f"Fallback extraction error: {str(e)}")
        #     raise

    def _generate_base_metadata(self, file_path: str) -> Dict:
        """Generate base metadata for a file with correct timestamp formats."""
        path = Path(file_path)
        stats = path.stat()

        try:
            # Calculate content hash
            hash_md5 = hashlib.md5()
            hash_sha256 = hashlib.sha256()

            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
                    hash_sha256.update(chunk)

            # Convert timestamps to milliseconds integers
            last_modified_ms = int(stats.st_mtime * 1000)
            created_at_ms = int(stats.st_ctime * 1000)

            return {
                "filename": path.name,
                "extension": path.suffix.lower(),
                "size": stats.st_size,  # Required field
                "last_modified": last_modified_ms,  # Integer timestamp in milliseconds
                "created_at": created_at_ms,  # Integer timestamp in milliseconds
                "content_hash": hash_sha256.hexdigest(),
                "file_path": str(path),
            }
        except Exception as e:
            logger.error(f"Error generating metadata for {file_path}: {str(e)}")
            raise

    def _post_process_content(self, content: str) -> str:
        """Clean and normalize extracted text."""
        if not content:
            return ""

        # Basic cleaning
        content = content.replace("\x00", "")  # Remove null bytes
        content = " ".join(content.split())  # Normalize whitespace
        content = content.strip()

        # Remove very short lines (often headers/footers)
        lines = [line for line in content.split("\n") if len(line.strip()) > 20]
        return "\n".join(lines)

    def _calculate_quality_score(self, content: str, metadata: Dict) -> float:
        """Calculate quality score for extracted content."""
        if not content:
            return 0.0

        factors = []

        # Content length score
        length_score = min(1.0, len(content) / 1000)
        factors.append(length_score)

        # Word variety score
        words = content.lower().split()
        if words:
            unique_ratio = len(set(words)) / len(words)
            factors.append(unique_ratio)

        # Format-specific scores
        if "encoding_confidence" in metadata:
            factors.append(metadata["encoding_confidence"])

        # Calculate final score
        return sum(factors) / len(factors) if factors else 0.0

    @staticmethod
    def _humanize_size(size_bytes: int) -> str:
        """Convert bytes to human readable format."""
        for unit in ["B", "KB", "MB", "GB"]:
            if size_bytes < 1024:
                return f"{size_bytes:.1f}{unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f}TB"


class TokenManager:
    def __init__(self, initial_token: str, server_url: str):
        self.token = initial_token
        self.server_url = server_url
        self.expiry_time = None
        self._lock = threading.Lock()
        logger.debug(f"TokenManager initialized with server URL: {server_url}")

    def validate_refresh_token(self) -> tuple[bool, str]:
        with self._lock:
            try:
                logger.debug("Attempting to validate/refresh token")
                response = requests.post(
                    f"{self.server_url}/auth/refresh",
                    headers={
                        "Authorization": f"Bearer {self.token}",
                        "Content-Type": "application/json",
                    },
                    timeout=50,
                )
                response.raise_for_status()
                data = response.json()

                self.token = JsonHelper.get_value(data, "access_token", self.token)
                self.is_valid = JsonHelper.get_value(data, "valid", False)

                if "expiry_time" in data:
                    self.expiry_time = datetime.fromisoformat(data["expiry_time"])
                    logger.debug(f"Token refreshed, expires at: {self.expiry_time}")

                return self.is_valid, self.token

            except Exception as e:
                logger.error(f"Token refresh failed: {str(e)}", exc_info=True)
                raise

    def get_valid_token(self) -> tuple[bool, str]:
        if not self.expiry_time or datetime.now() >= self.expiry_time:
            logger.error("Token expired, requesting refresh")
            return self.validate_refresh_token()
        return self.is_valid, self.token


class FolderEventHandler(FileSystemEventHandler):
    def __init__(self, watcher):
        self.watcher = watcher
        self.supported_extensions = SUPPORTED_EXTENSIONS
        super().__init__()
        logger.debug(
            "FolderEventHandler initialized with supported extensions: %s",
            self.supported_extensions,
        )

    def on_created(self, event):
        if (
            not event.is_directory
            and Path(event.src_path).suffix.lower() in self.supported_extensions
        ):
            logger.info(f"File created: {event.src_path}")
            self.watcher.process_file(event.src_path, "created")

    def on_modified(self, event):
        if (
            not event.is_directory
            and Path(event.src_path).suffix.lower() in self.supported_extensions
        ):
            logger.info(f"File modified: {event.src_path}")
            self.watcher.process_file(event.src_path, "modified")

    def on_deleted(self, event):
        if (
            not event.is_directory
            and Path(event.src_path).suffix.lower() in self.supported_extensions
        ):
            logger.info(f"File deleted: {event.src_path}")
            self.watcher.process_file(event.src_path, "deleted")


def load_env_file():
    # When running as PyInstaller bundle
    if getattr(sys, "frozen", False):
        application_path = sys._MEIPASS
    else:
        application_path = os.path.dirname(os.path.abspath(__file__))

    env_path = os.path.join(application_path, ".env")
    load_dotenv(env_path)


class FolderWatcher:
    def __init__(self):
        logger.info("Initializing FolderWatcher")

        # Load configuration
        load_env_file()

        self.server_url = os.getenv("SERVER_URL")
        self.connector_id = os.getenv("CONNECTOR_ID")
        self.port = int(os.getenv("PORT", "3000"))
        self.token = os.getenv("AUTH_TOKEN")
        self.folder_path = None

        if not all([self.server_url, self.connector_id, self.token]):
            logger.error("Missing required environment variables")
            raise ValueError("Missing required environment variables")

        # Initialize components
        self.token_manager = TokenManager(self.token, self.server_url)
        self.document_processor = DocumentProcessor()
        self.processed_files = {}  # Keep track of processed files and their metadata
        self.min_quality_threshold = 0.3

        # Watchdog components
        self.observer = None
        self.event_handler = None

        # Status tracking
        self.is_running = False
        self.shutdown_requested = False

        logger.debug(
            "FolderWatcher initialized with server URL: %s, connector ID: %s",
            self.server_url,
            self.connector_id,
        )

    def select_folder(self) -> str:
        logger.info("Opening folder selection dialog")
        root = tk.Tk()
        root.withdraw()
        folder = filedialog.askdirectory(title="Select Folder to Watch")
        if folder:
            logger.info(f"Selected folder: {folder}")
        else:
            logger.warning("No folder selected")
        return folder

    def perform_initial_scan(self):
        """Perform initial scan of all files in the watched folder."""
        logger.info(f"Starting initial scan of folder: {self.folder_path}")
        try:
            # Get list of all files in the directory and subdirectories
            for root, _, files in os.walk(self.folder_path):
                for filename in files:
                    file_path = os.path.join(root, filename)
                    if Path(file_path).suffix.lower() in SUPPORTED_EXTENSIONS:
                        try:
                            # Skip if file was already processed
                            if file_path in self.processed_files:
                                continue

                            logger.debug(f"Initial scan processing: {file_path}")
                            self.process_file(file_path, "created")

                        except Exception as e:
                            logger.error(
                                f"Error processing file {file_path} during initial scan: {str(e)}"
                            )
                            continue

            logger.info("Initial scan completed successfully")
            return True

        except Exception as e:
            logger.error(f"Failed to perform initial scan: {str(e)}")
            return False

    def initialize(self) -> bool:
        logger.info("Initializing watcher components")

        self.folder_path = self.select_folder()
        if not self.folder_path:
            logger.error("No folder selected, exiting")
            return False

        # Initial status checks
        if not self.check_auth_status():
            logger.error("Authentication check failed")
            return False

        if not self.check_connector_status():
            logger.error("Connector status check failed")
            return False

        # Perform initial scan
        if not self.perform_initial_scan():
            logger.error("Initial scan failed")
            return False

        # Initialize watchdog
        self.event_handler = FolderEventHandler(self)
        self.observer = Observer()
        self.observer.schedule(self.event_handler, self.folder_path, recursive=True)

        self.is_running = True
        logger.info("Watcher initialization completed successfully")
        return True

    def process_file(self, file_path: str, event_type: str):
        """Process a single file event."""
        try:
            logger.debug(f"Processing {event_type} event for file: {file_path}")

            if event_type == "deleted":
                self._handle_deleted_file(file_path)
            else:
                self._handle_active_file(file_path, event_type)

        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}", exc_info=True)
            self._update_file_status(file_path, "error", str(e))

    def start(self):
        """Start the folder watcher."""
        logger.info("Starting FolderWatcher")

        if not self.initialize():
            logger.error("Failed to initialize FolderWatcher")
            return

        self.observer.start()
        logger.info(f"Started watching folder: {self.folder_path}")

        try:
            while not self.shutdown_requested:
                # Periodic status checks
                if not self.check_auth_status() or not self.check_connector_status():
                    logger.error("Critical status check failed, stopping watcher")
                    break
                time.sleep(60)  # Status check interval

        except KeyboardInterrupt:
            logger.info("Received keyboard interrupt")
        finally:
            self.stop()

    def stop(self):
        """Stop the folder watcher."""
        logger.info("Stopping FolderWatcher")

        self.is_running = False
        self.shutdown_requested = True

        if self.observer:
            self.observer.stop()
            self.observer.join()

        logger.info("FolderWatcher stopped successfully")
        sys.exit(0)

    def check_auth_status(self) -> bool:
        """Check authentication status."""
        try:
            is_token_valid, auth_token = self.token_manager.get_valid_token()
            if not is_token_valid or not auth_token:
                logger.error("Authentication validation failed")
                self.shutdown_requested = True
                return False
            logger.debug("Authentication token is valid")
            return True
        except Exception as e:
            logger.error(f"Auth check failed: {str(e)}", exc_info=True)
            self.shutdown_requested = True
            return False

    def check_connector_status(self) -> bool:
        """Check if the connector is still active."""
        try:
            _, token = self.token_manager.get_valid_token()
            response = requests.get(
                f"{self.server_url}/connectors/folder/{self.connector_id}/status",
                headers={
                    "Authorization": f"Bearer {token}",
                    "Content-Type": "application/json",
                },
                timeout=300,
            )

            if response.status_code == 200:
                data = response.json()
                if data.get("active", True):
                    logger.debug(f"Connector {self.connector_id} is active")
                    return True

            logger.error("Connector is inactive")
            self.shutdown_requested = True
            return False

        except Exception as e:
            logger.error(f"Connector status check failed: {str(e)}", exc_info=True)
            self.shutdown_requested = True
            return False

    def _handle_deleted_file(self, file_path: str):
        """Handle deleted file event."""
        logger.debug(f"Handling deleted file: {file_path}")
        if file_path in self.processed_files:
            self._update_file_status(file_path, "deleted")
            self.processed_files.pop(file_path, None)

    def _handle_active_file(self, file_path: str, event_type: str):
        """Handle file creation or modification."""
        logger.debug(f"Processing active file: {file_path}")
        result = self.document_processor.process_file(file_path)

        if result.quality_score < self.min_quality_threshold:
            logger.warning(
                f"Low quality extraction for {file_path}: {result.quality_score}"
            )
            self._update_file_status(
                file_path, "error", f"Low quality extraction: {result.quality_score}"
            )
            return

        # Update processing history
        self.processed_files[file_path] = {
            "timestamp": time.time(),
            "metadata": result.metadata,
            "quality_score": result.quality_score,
        }

        # Send to backend
        payload = {
            "connector_id": self.connector_id,
            "event_type": event_type,
            "metadata": {
                **result.metadata,
                "extraction_quality": result.quality_score,
                "processing_timestamp": int(time.time() * 1000),
                "status": "active",
            },
            "content": result.content,
            "timestamp": int(time.time() * 1000),
        }

        self._send_to_backend(payload)

    def _update_file_status(
        self, file_path: str, status: str, error_message: Optional[str] = None
    ):
        """Update file status and send to server."""
        try:
            logger.debug(f"Updating file status: {file_path} -> {status}")
            metadata = self.processed_files.get(file_path, {}).get(
                "metadata", self.document_processor._generate_base_metadata(file_path)
            )

            metadata.update(
                {
                    "status": status,
                    "error_message": error_message,
                    "last_updated": int(time.time() * 1000),
                }
            )

            payload = {
                "connector_id": self.connector_id,
                "event_type": "status_update",
                "metadata": metadata,
                "timestamp": int(time.time() * 1000),
            }

            self._send_to_backend(payload)

        except Exception as e:
            logger.error(f"Failed to update file status: {str(e)}", exc_info=True)

    def _send_to_backend(self, payload: dict):
        """Send payload to backend with validation."""
        try:
            logger.debug(
                f"Sending payload to backend for file: {payload.get('metadata', {}).get('file_path')}"
            )

            _, token = self.token_manager.get_valid_token()
            response = requests.post(
                f"{self.server_url}/connectors/folder/watch",
                json=payload,
                headers={
                    "Authorization": f"Bearer {token}",
                    "Content-Type": "application/json",
                },
                timeout=3000,
            )
            response.raise_for_status()
            logger.debug("Successfully sent payload to backend")

        except Exception as e:
            logger.error(f"Failed to send to backend: {str(e)}", exc_info=True)
            raise


if __name__ == "__main__":
    try:
        watcher = FolderWatcher()
        watcher.start()
    except Exception as e:
        logger.error(f"Failed to start watcher: {e}")
        sys.exit(1)
