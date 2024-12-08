from pathlib import Path
import hashlib
from dataclasses import dataclass
from typing import Dict, Any, Optional, Set, TypeVar
import mimetypes
import chardet
import pdfplumber
from docx import Document
import magic
from PIL import Image
import pytesseract


@dataclass
class ExtractionResult:
    """Container for extraction results."""

    content: str
    metadata: Dict[str, Any]
    error: Optional[str] = None
    quality_score: float = 1.0


class DocumentProcessor:
    """Handles document processing and text extraction."""

    def __init__(self):
        # Initialize mime types
        mimetypes.init()
        # Configure minimum content quality thresholds
        self.min_content_length = 50
        self.min_quality_score = 0.3

    def process_file(self, file_path: str) -> ExtractionResult:
        """Process a file and extract text with metadata."""
        try:
            # Detect file type
            mime_type = magic.from_file(file_path, mime=True)

            # Get basic metadata first
            metadata = self._generate_base_metadata(file_path)

            # Extract text based on file type
            if mime_type == "application/pdf":
                result = self._extract_pdf(file_path)
            elif mime_type in [
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ]:
                result = self._extract_doc(file_path)
            elif mime_type.startswith("text/"):
                result = self._extract_text(file_path)
            elif mime_type.startswith("image/"):
                result = self._extract_image(file_path)
            else:
                # Fallback to textract for other formats
                result = self._extract_fallback(file_path)

            # Merge metadata
            metadata.update(result.metadata)

            # Post-process and validate content
            processed_content = self._post_process_content(result.content)
            quality_score = self._calculate_quality_score(
                processed_content, result.metadata
            )

            return ExtractionResult(
                content=processed_content,
                metadata=metadata,
                quality_score=quality_score,
            )

        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return ExtractionResult(
                content="",
                metadata=(
                    metadata
                    if "metadata" in locals()
                    else self._generate_base_metadata(file_path)
                ),
                error=str(e),
                quality_score=0.0,
            )

    def _extract_pdf(self, file_path: str) -> ExtractionResult:
        """Extract text from PDF files."""
        try:
            extracted_text = []
            metadata = {}

            with pdfplumber.open(file_path) as pdf:
                metadata.update(
                    {
                        "page_count": len(pdf.pages),
                        "pdf_info": pdf.metadata,
                    }
                )

                # Extract text from each page with layout preservation
                for i, page in enumerate(pdf.pages):
                    # Extract text with layout analysis
                    text = page.extract_text(layout=True)
                    if text:
                        extracted_text.append(text)

                    # Extract tables if present
                    tables = page.extract_tables()
                    if tables:
                        metadata[f"page_{i+1}_tables"] = len(tables)
                        # Convert tables to text representation
                        for table in tables:
                            table_text = "\n".join(
                                [" | ".join(row) for row in table if any(row)]
                            )
                            extracted_text.append(table_text)

            content = "\n\n".join(extracted_text)

            return ExtractionResult(
                content=content,
                metadata=metadata,
                quality_score=self._calculate_quality_score(content, metadata),
            )

        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise

    def _extract_doc(self, file_path: str) -> ExtractionResult:
        """Extract text from DOC/DOCX files."""
        try:
            doc = Document(file_path)
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
                "has_tables": bool(len(doc.tables)),
            }

            # Extract text with structure preservation
            content_parts = []

            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)

            # Process tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    content_parts.append("\n".join(table_text))

            content = "\n\n".join(content_parts)

            return ExtractionResult(
                content=content,
                metadata=metadata,
                quality_score=self._calculate_quality_score(content, metadata),
            )

        except Exception as e:
            logger.error(f"DOC extraction error: {str(e)}")
            raise

    def _extract_text(self, file_path: str) -> ExtractionResult:
        """Extract text from plain text files with encoding detection."""
        try:
            # Detect encoding
            with open(file_path, "rb") as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result["encoding"]
                confidence = result["confidence"]

            # Read with detected encoding
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()

            metadata = {
                "encoding": encoding,
                "encoding_confidence": confidence,
                "line_count": content.count("\n") + 1,
            }

            return ExtractionResult(
                content=content, metadata=metadata, quality_score=confidence
            )

        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            raise

    def _extract_image(self, file_path: str) -> ExtractionResult:
        """Extract text from images using OCR."""
        try:
            with Image.open(file_path) as img:
                metadata = {
                    "image_format": img.format,
                    "image_size": img.size,
                    "color_mode": img.mode,
                }

                # Perform OCR
                content = pytesseract.image_to_string(img)
                confidence = float(
                    pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)[
                        "conf"
                    ][0]
                )

                return ExtractionResult(
                    content=content, metadata=metadata, quality_score=confidence / 100.0
                )

        except Exception as e:
            logger.error(f"Image extraction error: {str(e)}")
            raise

    def _extract_fallback(self, file_path: str) -> ExtractionResult:
        """Fallback extraction using textract."""
        pass
        # try:
        #     content = textract.process(file_path).decode("utf-8")
        #     metadata = {"extraction_method": "textract"}

        #     return ExtractionResult(content=content, metadata=metadata)

        # except Exception as e:
        #     logger.error(f"Fallback extraction error: {str(e)}")
        #     raise

    def _generate_base_metadata(self, file_path: str) -> Dict:
        """Generate base metadata for a file with correct timestamp formats."""
        path = Path(file_path)
        stats = path.stat()

        try:
            # Calculate content hash
            hash_md5 = hashlib.md5()
            hash_sha256 = hashlib.sha256()

            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
                    hash_sha256.update(chunk)

            # Convert timestamps to milliseconds integers
            last_modified_ms = int(stats.st_mtime * 1000)
            created_at_ms = int(stats.st_ctime * 1000)

            return {
                "filename": path.name,
                "extension": path.suffix.lower(),
                "size": stats.st_size,  # Required field
                "last_modified": last_modified_ms,  # Integer timestamp in milliseconds
                "created_at": created_at_ms,  # Integer timestamp in milliseconds
                "content_hash": hash_sha256.hexdigest(),
                "file_path": str(path),
            }
        except Exception as e:
            logger.error(f"Error generating metadata for {file_path}: {str(e)}")
            raise

    def _post_process_content(self, content: str) -> str:
        """Clean and normalize extracted text."""
        if not content:
            return ""

        # Basic cleaning
        content = content.replace("\x00", "")  # Remove null bytes
        content = " ".join(content.split())  # Normalize whitespace
        content = content.strip()

        # Remove very short lines (often headers/footers)
        lines = [line for line in content.split("\n") if len(line.strip()) > 20]
        return "\n".join(lines)

    def _calculate_quality_score(self, content: str, metadata: Dict) -> float:
        """Calculate quality score for extracted content."""
        if not content:
            return 0.0

        factors = []

        # Content length score
        length_score = min(1.0, len(content) / 1000)
        factors.append(length_score)

        # Word variety score
        words = content.lower().split()
        if words:
            unique_ratio = len(set(words)) / len(words)
            factors.append(unique_ratio)

        # Format-specific scores
        if "encoding_confidence" in metadata:
            factors.append(metadata["encoding_confidence"])

        # Calculate final score
        return sum(factors) / len(factors) if factors else 0.0

    @staticmethod
    def _humanize_size(size_bytes: int) -> str:
        """Convert bytes to human readable format."""
        for unit in ["B", "KB", "MB", "GB"]:
            if size_bytes < 1024:
                return f"{size_bytes:.1f}{unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f}TB"
